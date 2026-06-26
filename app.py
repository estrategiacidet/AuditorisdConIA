from __future__ import annotations

import json
import os
import re
import time
from collections import OrderedDict
from pathlib import Path
from typing import Iterable

import openpyxl
import pypdf
import streamlit as st
from docx import Document
from google import genai
from google.genai import types

try:
    from docling.document_converter import DocumentConverter
except ModuleNotFoundError:
    DocumentConverter = None

try:
    from dotenv import load_dotenv
except ModuleNotFoundError:
    load_dotenv = None


BASE_DIR = Path(__file__).resolve().parent
EXCEL_TEMPLATE = BASE_DIR / "lista de verificación.xlsx"

if load_dotenv is not None:
    load_dotenv(BASE_DIR / ".env")

DEFAULT_API_KEY = os.getenv("GEMINI_API_KEY", "")
DEFAULT_MODELS = ("gemini-2.5-flash", "gemini-2.5-pro")
EVIDENCE_SOURCE_FOLDER_NAME = "Evidencias"
EVIDENCE_MARKDOWN_FOLDER_NAME = "Evidencias_Markdown"
DOCILING_SUPPORTED_EXTENSIONS = {
    ".bmp",
    ".csv",
    ".doc",
    ".docx",
    ".gif",
    ".htm",
    ".html",
    ".jpeg",
    ".jpg",
    ".md",
    ".pdf",
    ".png",
    ".pptx",
    ".tif",
    ".tiff",
    ".txt",
    ".webp",
    ".xls",
    ".xlsx",
}


INSTRUCCIONES_A1 = """
Eres un Consultor Senior Experto en Diagnóstico y Contexto Organizacional de CIDET. Tu único objetivo es recopilar, investigar y estructurar la información general y estratégica de la empresa asignada a partir de los documentos provistos y de su sitio web oficial.

Este diagnóstico NO es una auditoría formal; es una fase de reconocimiento preliminar para alimentar la revisión documental de un plan de auditoría.

LINEAMIENTOS ESTRICTOS DE COMPORTAMIENTO Y CALIDAD (CONTROL DE SESGOS):
1. Objetividad Absoluta: Está prohibido hacer juicios de valor, críticas, o declarar hallazgos, conformidades o incumplimientos normativos. Limítate a reportar datos objetivos de las fuentes.
2. Restricción de Evidencia: Bástate exclusivamente en los datos explícitos del documento de Cámara de Comercio, el organigrama, el Informe Guía suministrado y la URL web oficial que se te proporciona. Si un dato solicitado en la estructura no se encuentra en estas fuentes (por ejemplo, el organigrama detallado o el comité de SST), debes escribir textualmente: "Información no disponible en los documentos de entrada preliminares; se deberá validar en la auditoría de campo". Queda prohibido inventar o asumir datos.
3. Uso de la URL: Si se incluye una URL en los datos de entrada, navega en ella para extraer y validar la misión, visión, pilares estratégicos, presencia o menciones al sistema de gestión de la empresa.
4. Tono Corporativo: Utiliza una redacción neutral, técnica, clara y estrictamente profesional, mimetizando el lenguaje formal e institucional de CIDET.

PROCESO DE INVESTIGACIÓN:
FASE 0: INGESTA Y REVISIÓN DOCUMENTAL Y DIGITAL
- Analizar el documento de constitución (Cámara de Comercio).
- Analizar el organigrama de la empresa, si está disponible.
- Inspeccionar el contenido de la página web provista (si está disponible).
- Identificar cualquier referencia al organigrama, estructura o líderes de la empresa.

FASE 1: EXTRACCIÓN DE CONTEXTO
- Identificar organización, NIT, representantes legales, dirección, objeto social y códigos CIIU.
- Determinar las actividades principales y fortalezas estratégicas declaradas por la empresa en sus canales oficiales.
- Identificar menciones a actores clave, comités de SST (COPASST) o vigencia de acreditaciones bajo ISO 45001.

FASE 2: CONSOLIDACIÓN DE INFORMACIÓN
Estructurar los datos extraídos en los siguientes ejes antes de redactar: Tipo de empresa, años de trayectoria, actividades principales, actores clave, estado preliminar de SST y fortalezas estratégicas.

FASE 3: GENERACIÓN DE INFORME (ESTRUCTURA OBLIGATORIA DE SALIDA)
Debes estructurar y entregar el informe siguiendo exactamente este formato de salida. No agregues introducciones ni conclusiones fuera de los bloques delimitados:

[INICIO_CONTEXTO]

# INFORME CONTEXTO ORGANIZACIONAL
## Empresa: [Insertar Nombre de la Empresa completo y NIT si está disponible]
## Fecha de Generación: 24 de junio de 2026

### 1. RESUMEN EJECUTIVO
[Redacta un resumen analítico y fluido de un máximo de dos párrafos que sintetice la naturaleza de la empresa analizada, su propósito principal y el alcance del reconocimiento preliminar realizado para CIDET].

### 2. INFORMACIÓN GENERAL Y MARCO OPERATIVO
- **Razón Social Completa:** [Nombre]
- **NIT / Identificación Legal:** [Número]
- **Actividad Económica Principal (CIIU):** [Indicar actividad y código si aparece]
- **Años de Trayectoria / Constitución:** [Fecha o años desde su constitución legal]
- **Representante Legal / Actores Clave:** [Nombres de cargos directivos identificados]
- **Estructura Organizacional preliminar:** [Describir brevemente cómo se organiza según el texto/web o reportar si no está disponible]
- **Documentos ausentes o no localizados:** [Si el organigrama o la Cámara de Comercio no se encuentran, detállalo de forma explícita]

### 3. CONTEXTO ESTRATÉGICO Y COMPONENTE EN SST
- **Actividades Principales del Negocio:** [Detallar las operaciones críticas de la organización recopiladas de los documentos y su sitio web]
- **Fortalezas Estratégicas Identificadas:** [Mencionar las ventajas operativas, misión, visión o ventajas competitivas explícitas en las fuentes]
- **Comité de SST / Estado ISO 45001:** [Reportar lo hallado en los documentos o en el sitio web sobre comités o certificaciones de seguridad, o declarar su ausencia documental para validación en campo]

[FIN_CONTEXTO]
"""


def normalize_path(raw_value: str) -> Path:
    cleaned = raw_value.strip().strip('"').strip("'")
    return Path(cleaned).expanduser().resolve()


def is_retryable_genai_error(exc: Exception) -> bool:
    message = str(exc).lower()
    retryable_markers = (
        "503",
        "unavailable",
        "high demand",
        "service unavailable",
        "temporarily unavailable",
        "try again later",
    )
    return any(marker in message for marker in retryable_markers)


def generate_content_with_retry(
    client: genai.Client,
    models: tuple[str, ...],
    *,
    contents: str,
    config: types.GenerateContentConfig,
    label: str,
    max_attempts: int = 3,
    base_delay_seconds: float = 2.0,
):
    last_error: Exception | None = None

    for model_name in models:
        delay = base_delay_seconds
        for attempt in range(1, max_attempts + 1):
            try:
                return client.models.generate_content(
                    model=model_name,
                    contents=contents,
                    config=config,
                )
            except Exception as exc:
                last_error = exc
                if not is_retryable_genai_error(exc) or attempt == max_attempts:
                    break
                time.sleep(delay)
                delay = min(delay * 2, 10.0)

    tried_models = ", ".join(models)
    raise RuntimeError(
        f"{label} no pudo completarse porque Gemini devolvió un error temporal. "
        f"Se reintentó la solicitud y se probaron estos modelos: {tried_models}. "
        f"Detalle original: {last_error}"
    ) from last_error


def iter_pdf_text(file_paths: Iterable[Path]) -> str:
    chunks: list[str] = []
    for pdf_path in file_paths:
        try:
            reader = pypdf.PdfReader(str(pdf_path))
            for page in reader.pages:
                text = page.extract_text() or ""
                if text:
                    chunks.append(text)
        except Exception:
            continue
    return "\n".join(chunks)


def is_ignored_generated_file(path: Path) -> bool:
    lower_name = path.name.lower()
    ignored_tokens = (
        "contexto_organizacional",
        "informe_final_auditoria_iso45001",
        "lista_de_verificacion_llena",
        "lista de verificaci",
    )
    return any(token in lower_name for token in ignored_tokens)


def iter_supported_text_files(base_dir: Path, exclude_folder_names: set[str] | None = None) -> list[Path]:
    exclude_folder_names = exclude_folder_names or set()
    files: list[Path] = []
    for candidate in base_dir.rglob("*"):
        if not candidate.is_file():
            continue
        if candidate.suffix.lower() not in {".pdf", ".docx", ".md", ".txt"}:
            continue
        if is_ignored_generated_file(candidate):
            continue
        if any(part.lower() in exclude_folder_names for part in candidate.parts):
            continue
        files.append(candidate)
    return files


def read_docx_text(file_path: Path) -> str:
    try:
        document = Document(str(file_path))
    except Exception:
        return ""
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    return "\n".join(paragraphs)


def read_text_from_file(file_path: Path) -> str:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return iter_pdf_text([file_path]).strip()
    if suffix == ".docx":
        return read_docx_text(file_path).strip()
    try:
        return file_path.read_text(encoding="utf-8", errors="ignore").strip()
    except Exception:
        return ""


def find_evidence_folder(base_dir: Path) -> Path:
    for candidate in base_dir.rglob("*"):
        if candidate.is_dir() and "evidenc" in candidate.name.lower() and "markdown" not in candidate.name.lower():
            return candidate
    return base_dir / EVIDENCE_SOURCE_FOLDER_NAME


def get_evidence_markdown_folder(base_dir: Path) -> Path:
    return base_dir / EVIDENCE_MARKDOWN_FOLDER_NAME


def iter_evidence_source_files(folder: Path) -> list[Path]:
    files: list[Path] = []
    if not folder.exists():
        return files
    for candidate in folder.rglob("*"):
        if not candidate.is_file():
            continue
        if candidate.suffix.lower() not in DOCILING_SUPPORTED_EXTENSIONS:
            continue
        files.append(candidate)
    return files


def build_markdown_output_path(source_file: Path, source_root: Path, markdown_root: Path) -> Path:
    relative_path = source_file.relative_to(source_root)
    return markdown_root / relative_path.with_suffix(".md")


def convert_evidences_to_markdown(
    source_folder: Path,
    markdown_folder: Path,
) -> dict[str, int | list[str] | str]:
    source_files = iter_evidence_source_files(source_folder)
    result: dict[str, int | list[str] | str] = {
        "source_folder": str(source_folder),
        "markdown_folder": str(markdown_folder),
        "processed": 0,
        "converted": 0,
        "skipped": 0,
        "failed": [],
    }

    if not source_folder.exists():
        return result

    if DocumentConverter is None:
        raise RuntimeError(
            "Docling no está instalado en el entorno actual. "
            "Instala la dependencia `docling` para convertir las evidencias a Markdown."
        )

    markdown_folder.mkdir(parents=True, exist_ok=True)
    converter = DocumentConverter()

    for source_file in source_files:
        result["processed"] += 1
        output_file = build_markdown_output_path(source_file, source_folder, markdown_folder)
        output_file.parent.mkdir(parents=True, exist_ok=True)

        try:
            if output_file.exists() and output_file.stat().st_mtime >= source_file.stat().st_mtime:
                result["skipped"] += 1
                continue

            conversion_result = converter.convert(str(source_file))
            markdown_content = conversion_result.document.export_to_markdown().strip()
            if not markdown_content:
                markdown_content = f"# {source_file.name}\n\nNo se pudo extraer contenido legible con Docling."

            output_file.write_text(markdown_content, encoding="utf-8")
            result["converted"] += 1
        except Exception as exc:
            failed_entries = result["failed"]
            if isinstance(failed_entries, list):
                failed_entries.append(f"{source_file.name}: {exc}")

    return result


def collect_document_bundle(base_dir: Path) -> dict[str, dict[str, list[Path] | str]]:
    bundle: dict[str, dict[str, list[Path] | str]] = {
        "camara": {"paths": [], "text": ""},
        "organigrama": {"paths": [], "text": ""},
        "guia": {"paths": [], "text": ""},
    }
    for file_path in iter_supported_text_files(
        base_dir,
        exclude_folder_names={"evidencias", "evidencias_markdown"},
    ):
        name_lower = file_path.name.lower()
        content = read_text_from_file(file_path)
        content_lower = content.lower()

        if any(keyword in name_lower or keyword in content_lower for keyword in ("cámara de comercio", "camara de comercio", "constitución", "constitucion", "existencia y representación", "existencia y representacion", "representación legal", "representacion legal", "cámara", "camara")):
            bundle["camara"]["paths"].append(file_path)
            if content:
                bundle["camara"]["text"] = f"{bundle['camara']['text']}\n{content}".strip()

        if any(keyword in name_lower or keyword in content_lower for keyword in ("organigrama", "estructura organizacional", "estructura orgánica", "estructura organica")):
            bundle["organigrama"]["paths"].append(file_path)
            if content:
                bundle["organigrama"]["text"] = f"{bundle['organigrama']['text']}\n{content}".strip()

        if any(keyword in name_lower or keyword in content_lower for keyword in ("informe", "guía", "guia")):
            bundle["guia"]["paths"].append(file_path)
            if content:
                bundle["guia"]["text"] = f"{bundle['guia']['text']}\n{content}".strip()

    return bundle


def collect_evidence_text(folder: Path) -> str:
    if not folder.exists():
        return ""
    supported_files = iter_supported_text_files(folder)
    texts: list[str] = []
    for file_path in supported_files:
        text = read_text_from_file(file_path)
        if not text:
            continue
        texts.append(f"# Archivo: {file_path.name}\n\n{text}")
    return "\n\n".join(texts)


def build_context_package(base_dir: Path, url_empresa: str) -> tuple[str, dict[str, list[str]]]:
    bundle = collect_document_bundle(base_dir)
    found_docs = {key: [path.name for path in value["paths"]] for key, value in bundle.items()}

    texto_camara = bundle["camara"]["text"].strip()
    texto_organigrama = bundle["organigrama"]["text"].strip()
    texto_guia = bundle["guia"]["text"].strip()

    paquete = (
        f"[CÁMARA DE COMERCIO / CONSTITUCIÓN]:\n{texto_camara or 'No localizado'}\n\n"
        f"[ORGANIGRAMA]:\n{texto_organigrama or 'No localizado'}\n\n"
        f"[EJEMPLO INFORME]:\n{texto_guia or 'No localizado'}"
    )
    if url_empresa.strip():
        paquete += f"\n\n[SITIO WEB OFICIAL A ANALIZAR]: {url_empresa.strip()}"

    faltantes: list[str] = []
    if not bundle["camara"]["paths"]:
        faltantes.append("Documento de Cámara de Comercio / constitución")
    if not bundle["organigrama"]["paths"]:
        faltantes.append("Organigrama")
    if not bundle["guia"]["paths"]:
        faltantes.append("Informe guía")
    if faltantes:
        paquete += (
            "\n\n[AVISO DE COBERTURA DOCUMENTAL]: No se localizaron los siguientes documentos en la carpeta del cliente: "
            + ", ".join(faltantes)
            + ". Debes notificarlo explícitamente en el informe de contexto y continuar con el análisis del resto de fuentes."
        )

    return paquete, found_docs


def read_context_files(base_dir: Path) -> tuple[str, str]:
    texto_camara = ""
    texto_guia = ""
    pdfs = [p for p in base_dir.iterdir() if p.is_file() and p.suffix.lower() == ".pdf"]
    for pdf in pdfs:
        name = pdf.name.lower()
        if "camara" in name:
            texto_camara += iter_pdf_text([pdf]) + "\n"
        elif "informe" in name:
            texto_guia += iter_pdf_text([pdf]) + "\n"
    return texto_camara.strip(), texto_guia.strip()


def read_excel_numerals(template_path: Path) -> list[str]:
    workbook = openpyxl.load_workbook(template_path, data_only=True)
    try:
        sheet = workbook["lista de verificación"] if "lista de verificación" in workbook.sheetnames else workbook.active
        values: list[str] = []
        for row in range(2, 149):
            cell_value = sheet.cell(row=row, column=4).value
            if cell_value not in (None, ""):
                values.append(str(cell_value).strip())
        return values
    finally:
        workbook.close()


def read_evidence_text(folder: Path) -> str:
    if not folder.exists():
        return ""
    pdfs = [p for p in folder.iterdir() if p.is_file() and p.suffix.lower() == ".pdf"]
    return iter_pdf_text(pdfs)


def read_checklist_rows(template_path: Path) -> list[dict[str, str]]:
    workbook = openpyxl.load_workbook(template_path, data_only=True)
    try:
        sheet = workbook["lista de verificaciÃ³n"] if "lista de verificaciÃ³n" in workbook.sheetnames else workbook.active
        rows: list[dict[str, str]] = []
        for row_index in range(2, 149):
            row_id = sheet.cell(row=row_index, column=1).value
            if row_id in (None, ""):
                continue
            rows.append(
                {
                    "excel_row": str(row_index),
                    "id": str(row_id).strip(),
                    "ciclo": str(sheet.cell(row=row_index, column=2).value or "").strip(),
                    "capitulo": str(sheet.cell(row=row_index, column=3).value or "").strip(),
                    "clausula": str(sheet.cell(row=row_index, column=4).value or "").strip(),
                    "tema": str(sheet.cell(row=row_index, column=5).value or "").strip(),
                    "requisito": str(sheet.cell(row=row_index, column=6).value or "").strip(),
                    "pregunta": str(sheet.cell(row=row_index, column=7).value or "").strip(),
                    "evidencia_esperada": str(sheet.cell(row=row_index, column=8).value or "").strip(),
                    "metodo": str(sheet.cell(row=row_index, column=10).value or "").strip(),
                    "roles": str(sheet.cell(row=row_index, column=11).value or "").strip(),
                    "referencia_base": str(sheet.cell(row=row_index, column=22).value or "").strip(),
                }
            )
        return rows
    finally:
        workbook.close()


def format_checklist_rows(rows: list[dict[str, str]]) -> str:
    lines: list[str] = []
    for row in rows:
        lines.append(
            " | ".join(
                [
                    row["id"],
                    row["capitulo"],
                    row["clausula"],
                    row["tema"],
                    row["pregunta"],
                ]
            )
        )
    return "\n".join(lines)


def parse_json_response(raw_text: str) -> dict:
    cleaned = raw_text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\s*```$", "", cleaned)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        start = cleaned.find("{")
        end = cleaned.rfind("}")
        if start != -1 and end != -1 and end > start:
            candidate = cleaned[start : end + 1]
            return json.loads(candidate)
        raise


def normalize_status(value: str) -> str:
    normalized = value.strip().lower()
    mapping = {
        "cumple": "Cumple",
        "cumple parcial": "Cumple parcial",
        "parcial": "Cumple parcial",
        "no cumple": "No cumple",
        "no aplica": "No aplica",
        "no evaluado": "No evaluado",
    }
    return mapping.get(normalized, value.strip() or "No evaluado")


def normalize_score(value: object, status: str) -> object:
    if value in (None, ""):
        status_lower = status.lower()
        if status_lower == "cumple":
            return 1
        if status_lower == "cumple parcial":
            return 0.5
        if status_lower == "no cumple":
            return 0
        return ""
    if isinstance(value, (int, float)):
        return value
    if isinstance(value, str):
        candidate = value.strip().replace(",", ".")
        try:
            numeric_value = float(candidate)
            if numeric_value.is_integer():
                return int(numeric_value)
            return numeric_value
        except ValueError:
            return value
    return value


def fill_checklist_workbook(template_path: Path, output_path: Path, row_results: list[dict]) -> None:
    workbook = openpyxl.load_workbook(template_path)
    try:
        sheet = workbook["lista de verificaciÃ³n"] if "lista de verificaciÃ³n" in workbook.sheetnames else workbook.active
        by_id = {str(row.get("id", "")).strip(): row for row in row_results if row.get("id")}
        for row_index in range(2, 149):
            row_id = sheet.cell(row=row_index, column=1).value
            if not row_id:
                continue
            payload = by_id.get(str(row_id).strip())
            if not payload:
                continue
            status = normalize_status(str(payload.get("estado", "No evaluado")))
            sheet.cell(row=row_index, column=12).value = status
            sheet.cell(row=row_index, column=13).value = normalize_score(payload.get("puntaje"), status)
            sheet.cell(row=row_index, column=14).value = str(payload.get("tipo_hallazgo", "")).strip()
            sheet.cell(row=row_index, column=15).value = str(payload.get("criticidad", "")).strip()
            sheet.cell(row=row_index, column=16).value = str(payload.get("evidencia_encontrada", "")).strip()
            sheet.cell(row=row_index, column=17).value = str(payload.get("analisis_hallazgo", "")).strip()
            sheet.cell(row=row_index, column=18).value = str(payload.get("accion_recomendada", "")).strip()
            sheet.cell(row=row_index, column=19).value = str(payload.get("responsable", "")).strip()
            sheet.cell(row=row_index, column=20).value = str(payload.get("fecha_compromiso", "")).strip()
            sheet.cell(row=row_index, column=21).value = str(payload.get("requiere_accion_correctiva", "")).strip()
            sheet.cell(row=row_index, column=22).value = str(payload.get("referencia_norma", "")).strip()
            sheet.cell(row=row_index, column=23).value = str(payload.get("fuente", "")).strip()
            sheet.cell(row=row_index, column=24).value = str(payload.get("notas", "")).strip()
        workbook.save(str(output_path))
    finally:
        workbook.close()


def build_chapter_summary(rows_payload: list[dict]) -> str:
    grouped: OrderedDict[str, list[dict]] = OrderedDict()
    for row in rows_payload:
        chapter = str(row.get("capitulo", "")).strip() or "Sin capítulo"
        grouped.setdefault(chapter, []).append(row)
    summary_lines: list[str] = []
    for chapter, items in grouped.items():
        counts = {
            "Cumple": 0,
            "Cumple parcial": 0,
            "No cumple": 0,
            "No aplica": 0,
            "No evaluado": 0,
        }
        for item in items:
            status = normalize_status(str(item.get("estado", "No evaluado")))
            counts[status] = counts.get(status, 0) + 1
        summary_lines.append(
            f"- {chapter}: Cumple={counts['Cumple']}, Parcial={counts['Cumple parcial']}, No cumple={counts['No cumple']}, No aplica={counts['No aplica']}, No evaluado={counts['No evaluado']}"
        )
    return "\n".join(summary_lines)


def build_final_report(rows_payload: list[dict], contexto_previo: str) -> str:
    total_rows = len(rows_payload)
    counts = {
        "Cumple": 0,
        "Cumple parcial": 0,
        "No cumple": 0,
        "No aplica": 0,
        "No evaluado": 0,
    }
    for row in rows_payload:
        status = normalize_status(str(row.get("estado", "No evaluado")))
        counts[status] = counts.get(status, 0) + 1

    chapter_summary = build_chapter_summary(rows_payload)
    conformes = [row for row in rows_payload if normalize_status(str(row.get("estado", ""))) == "Cumple"]
    oportunidades = [
        row
        for row in rows_payload
        if normalize_status(str(row.get("estado", ""))) in {"Cumple parcial", "No cumple", "No evaluado"}
    ]

    fortalezas_lines = []
    for row in conformes[:8]:
        fortalezas_lines.append(
            f"- {row.get('id', 'N/A')}: {str(row.get('evidencia_encontrada', '')).strip() or 'Evidencia documentada en la matriz.'}"
        )
    if not fortalezas_lines:
        fortalezas_lines.append("- No se identificaron conformidades concluyentes en la matriz procesada.")

    mejora_lines = []
    for row in oportunidades[:10]:
        mejora_lines.append(
            f"- {row.get('id', 'N/A')}: {str(row.get('analisis_hallazgo', '')).strip() or 'Revisar la evidencia disponible y completar la validación documental.'}"
        )
    if not mejora_lines:
        mejora_lines.append("- No se identificaron oportunidades de mejora relevantes con la evidencia disponible.")

    contexto_resumen = contexto_previo.strip().splitlines()
    contexto_excerpt = "\n".join(contexto_resumen[:8]).strip()
    if not contexto_excerpt:
        contexto_excerpt = "No se dispuso de un contexto organizacional legible para la redacción del informe."

    return "\n".join(
        [
            "# INFORME EJECUTIVO FINAL DE AUDITORÍA ISO 45001",
            "## CIDET",
            "",
            "### 1. Alcance y fuentes",
            "El presente informe consolida la lectura del contexto organizacional generado por el Agente 1, la estructura de la matriz de verificación y las evidencias físicas disponibles en la carpeta del cliente.",
            "",
            "### 2. Síntesis del contexto",
            contexto_excerpt,
            "",
            "### 3. Resumen de resultados de la matriz",
            f"- Total de registros evaluados: {total_rows}",
            f"- Cumple: {counts['Cumple']}",
            f"- Cumple parcial: {counts['Cumple parcial']}",
            f"- No cumple: {counts['No cumple']}",
            f"- No aplica: {counts['No aplica']}",
            f"- No evaluado: {counts['No evaluado']}",
            "",
            "### 4. Resumen por capítulos",
            chapter_summary or "- No fue posible consolidar capítulos.",
            "",
            "### 5. Fortalezas y hallazgos conformes",
            "\n".join(fortalezas_lines),
            "",
            "### 6. Aspectos por mejorar",
            "\n".join(mejora_lines),
            "",
            "### 7. Conclusión",
            "Con base en la matriz procesada, el informe refleja el estado documental disponible al momento del análisis. Las conclusiones deben complementarse con validación de campo cuando existan registros marcados como no evaluados o con evidencia insuficiente.",
        ]
    )


def build_local_context_report(
    base_dir: Path,
    url_empresa: str,
    documentos_detectados: dict[str, list[str]],
    paquete_contexto_base: str,
) -> str:
    camara = ", ".join(documentos_detectados.get("camara", [])) or "No localizado"
    organigrama = ", ".join(documentos_detectados.get("organigrama", [])) or "No localizado"
    guia = ", ".join(documentos_detectados.get("guia", [])) or "No localizado"
    url_texto = url_empresa.strip() or "No se proporcionó URL"

    return "\n".join(
        [
            "# REPORTE DE CONTEXTO ORGANIZACIONAL",
            "## Generación de respaldo local",
            "",
            "### 1. Ruta analizada",
            str(base_dir),
            "",
            "### 2. URL asociada",
            url_texto,
            "",
            "### 3. Documentos detectados",
            f"- Cámara de Comercio / constitución: {camara}",
            f"- Organigrama: {organigrama}",
            f"- Informe guía: {guia}",
            "",
            "### 4. Insumo consolidado",
            paquete_contexto_base.strip(),
        ]
    )


def build_agent2_local_report(contexto_previo: str, texto_evidencias: str) -> str:
    contexto_lineas = [line.strip() for line in contexto_previo.splitlines() if line.strip()]
    evidencia_lineas = [line.strip() for line in texto_evidencias.splitlines() if line.strip()]
    contexto_excerpt = "\n".join(contexto_lineas[:8]).strip() or "No se dispuso de un contexto legible."
    evidencia_excerpt = "\n".join(evidencia_lineas[:12]).strip() or "No se localizaron evidencias legibles en la carpeta del cliente."

    return "\n".join(
        [
            "# INFORME FINAL DE HALLAZGOS",
            "",
            "## 1. Resumen ejecutivo",
            "Con la información disponible en el contexto organizacional y en las evidencias físicas localizadas, no fue posible realizar una evaluación exhaustiva de todos los componentes del sistema. El análisis se limita a la trazabilidad documental observada y a la cobertura real de las fuentes encontradas.",
            "",
            "## 2. Contexto analizado",
            contexto_excerpt,
            "",
            "## 3. Evidencias revisadas",
            evidencia_excerpt,
            "",
            "## 4. Hallazgos preliminares",
            "- No se identificó una base documental suficiente para sostener una evaluación integral de todos los requisitos revisados.",
            "- La lectura de las evidencias debe complementarse con validación de campo cuando existan vacíos, documentos incompletos o ausencia de trazabilidad.",
            "",
            "## 5. Conclusión",
            "El resultado preliminar confirma la necesidad de profundizar la validación documental y operativa antes de emitir conclusiones definitivas sobre conformidad o no conformidad.",
        ]
    )


def save_docx_report(title: str, body: str, output_path: Path) -> None:
    doc = Document()
    doc.add_heading(title, level=1)
    for line in body.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            continue
        doc.add_paragraph(stripped)
    doc.save(str(output_path))


st.set_page_config(page_title="CIDET - Asistente de Auditoría IA", page_icon="⚡", layout="wide")

st.title("⚡ CIDET - Asistente de Auditoría Asistida con IA.")
st.markdown("Optimización y automatización de diagnóstico organizacional y evaluación ISO 45001")
st.write("---")

with st.sidebar:
    st.header("🔑 Configuración de Seguridad")
    api_key_usuario = st.text_input(
        "Gemini API Key",
        value=DEFAULT_API_KEY,
        type="password",
        help="También puedes definir la variable de entorno GEMINI_API_KEY.",
    )
    st.write("---")
    st.markdown("**Desarrollado para el auditor en sistemas de gestión - CIDET**")

st.header("📁 1. Información del cliente")

ruta_base = st.text_input(
    "Pega aquí la ruta de la auditoría de la empresa que vas a procesar hoy:",
    placeholder=r"C:\Users\usuario\OneDrive - CIDET\...",
)

url_empresa = st.text_input(
    "Pega aquí el enlace de la página web de la empresa (opcional):",
    placeholder="https://www.empresa.com",
)

if ruta_base:
    ruta_base_limpia = normalize_path(ruta_base)
    if not ruta_base_limpia.exists():
        st.error(f"No se encontró la carpeta indicada: `{ruta_base_limpia}`")
        st.stop()

    ruta_evidencias = find_evidence_folder(ruta_base_limpia)
    ruta_evidencias_markdown = get_evidence_markdown_folder(ruta_base_limpia)
    paquete_contexto_base, documentos_detectados = build_context_package(ruta_base_limpia, url_empresa)
    documentos_faltantes = []
    if not documentos_detectados.get("camara"):
        documentos_faltantes.append("Cámara de Comercio / constitución")
    if not documentos_detectados.get("organigrama"):
        documentos_faltantes.append("Organigrama")
    if not documentos_detectados.get("guia"):
        documentos_faltantes.append("Informe guía")

    st.success("📍 Conectado exitosamente a la carpeta del cliente.")
    if ruta_evidencias.exists():
        st.caption(f"Subcarpeta de evidencias detectada en: `{ruta_evidencias}`")
        st.caption(f"Los Markdown normalizados se guardarán en: `{ruta_evidencias_markdown}`")
    else:
        st.warning("No se encontró automáticamente la subcarpeta de evidencias; se intentará usar la ruta esperada `Evidencias`.")
    if documentos_faltantes:
        st.warning("Documentos no localizados en la carpeta del cliente: " + ", ".join(documentos_faltantes))

    tab1, tab2 = st.tabs(["🕵️ Agente 1: Investigador de Contexto", "📋 Agente 2: Hallazgos e Informe Final"])

    with tab1:
        st.subheader("An?lisis de C?mara de Comercio, Organigrama, Sitio Web e Informe Gu?a")

        if st.button("?? Ejecutar Agente Investigador", key="btn_agente1"):
            if not api_key_usuario.strip():
                st.error("Debes configurar tu Gemini API Key para ejecutar esta fase.")
            else:
                with st.spinner("El Agente 1 est? analizando los documentos..."):
                    try:
                        client = genai.Client(api_key=api_key_usuario.strip())

                        response = generate_content_with_retry(
                            client,
                            DEFAULT_MODELS,
                            contents=paquete_contexto_base,
                            config=types.GenerateContentConfig(
                                system_instruction=INSTRUCCIONES_A1,
                                temperature=0.2,
                            ),
                            label="El Agente 1",
                        )

                        ruta_out_word = ruta_base_limpia / "Contexto_Organizacional.docx"
                        ruta_out_txt = ruta_base_limpia / "Contexto_Organizacional.txt"

                        save_docx_report(
                            "Reporte de Contexto Organizacional y Pre-Diagn?stico",
                            response.text or "",
                            ruta_out_word,
                        )
                        ruta_out_txt.write_text(response.text or "", encoding="utf-8")

                        st.success("? ?Documento Word de Contexto generado con ?xito!")
                        st.text_area("Vista previa del Contexto:", response.text or "", height=300)
                    except Exception as exc:
                        if isinstance(exc, RuntimeError) and "Gemini devolvió un error temporal" in str(exc):
                            fallback_text = build_local_context_report(
                                ruta_base_limpia,
                                url_empresa,
                                documentos_detectados,
                                paquete_contexto_base,
                            )
                            ruta_out_word = ruta_base_limpia / "Contexto_Organizacional.docx"
                            ruta_out_txt = ruta_base_limpia / "Contexto_Organizacional.txt"
                            save_docx_report(
                                "Reporte de Contexto Organizacional y Pre-Diagnóstico",
                                fallback_text,
                                ruta_out_word,
                            )
                            ruta_out_txt.write_text(fallback_text, encoding="utf-8")
                            st.warning(
                                "El Agente 1 no logró usar Gemini por saturación temporal, pero la aplicación generó un contexto de respaldo local para no bloquear el flujo."
                            )
                            st.text_area("Vista previa del Contexto de respaldo:", fallback_text, height=300)
                        else:
                            st.error(f"Ocurrió un error con el Agente 1: {exc}")

    with tab2:
        st.subheader("Generación de hallazgos e informe final")
        st.write("Esta fase leerá el contexto generado por el Agente 1 y las evidencias del cliente para producir hallazgos y redactar el informe final, sin depender de la lista de verificación.")

        if st.button("?? Generar hallazgos e informe final", key="btn_agente2"):
            if not api_key_usuario.strip():
                st.error("Debes configurar tu Gemini API Key para ejecutar esta fase.")
            else:
                archivo_contexto = ruta_base_limpia / "Contexto_Organizacional.txt"
                if not archivo_contexto.exists():
                    st.error("No se encontr? el Contexto Organizacional. Debes ejecutar primero el Agente 1.")
                else:
                    with st.spinner("Leyendo el contexto, evaluando evidencias y generando hallazgos e informe final..."):
                        try:
                            client = genai.Client(api_key=api_key_usuario.strip())

                            contexto_previo = archivo_contexto.read_text(encoding="utf-8")
                            conversion_summary = convert_evidences_to_markdown(
                                ruta_evidencias,
                                ruta_evidencias_markdown,
                            )
                            texto_evidencias = collect_evidence_text(ruta_evidencias_markdown)
                            if not texto_evidencias.strip():
                                raise RuntimeError(
                                    "No se obtuvo contenido Markdown legible desde la carpeta de evidencias."
                                )

                            prompt_informe_completo = f"""
Actúas como un Lead Auditor Senior de Sistemas de Gestión en CIDET.
Debes devolver únicamente un informe en texto plano, sin bloques de código ni JSON.
Tu prioridad es identificar hallazgos, observaciones y no conformidades con base exclusivamente en el contexto organizacional generado por el Agente 1 y en las evidencias físicas del cliente.
No leas ni utilices ninguna lista de verificación, matriz Excel o numerales obligatorios.

INSUMO CLAVE 1 - Contexto Organizacional (Agente 1):
{contexto_previo}

INSUMO CLAVE 2 - Evidencias físicas del cliente:
{texto_evidencias}

Estructura obligatoria del informe:
1. Título.
2. Resumen ejecutivo de máximo dos párrafos.
3. Hallazgos preliminares, en viñetas, con redacción técnica y objetiva.
4. Conclusión.

Reglas:
- Usa únicamente la información disponible en las fuentes proporcionadas.
- Si una evidencia no permite concluir, indícalo explícitamente como limitación.
- No inventes datos no presentes en las fuentes.
"""

                            informe_res = generate_content_with_retry(
                                client,
                                DEFAULT_MODELS,
                                contents=prompt_informe_completo,
                                config=types.GenerateContentConfig(
                                    temperature=0.2,
                                ),
                                label="El Agente 2",
                            )

                            report_text = (informe_res.text or "").strip()
                            if not report_text:
                                report_text = build_agent2_local_report(contexto_previo, texto_evidencias)

                            ruta_word_final = ruta_base_limpia / "Informe_Final_Auditoria_ISO45001.docx"
                            save_docx_report(
                                "Informe Ejecutivo Final de Auditor?a ISO 45001 - CIDET",
                                report_text,
                                ruta_word_final,
                            )

                            if conversion_summary["processed"]:
                                st.info(
                                    "Conversión a Markdown completada: "
                                    f"{conversion_summary['converted']} convertidos, "
                                    f"{conversion_summary['skipped']} reutilizados y "
                                    f"{len(conversion_summary['failed'])} con error."
                                )
                            failed_entries = conversion_summary["failed"]
                            if isinstance(failed_entries, list) and failed_entries:
                                st.warning(
                                    "Algunas evidencias no pudieron convertirse con Docling: "
                                    + " | ".join(failed_entries[:5])
                                )

                            st.success("?? ?Informe ejecutivo en Word generado con ?xito!")
                            st.info("Por ahora se omiti? la generaci?n del Excel para no bloquear la producci?n del informe final.")
                            st.markdown(f"?? **Informe guardado en:** `{ruta_word_final}`")
                            st.subheader("?? Vista Previa del Informe Redactado:")
                            st.text_area("Contenido:", report_text, height=400)
                        except Exception as exc:
                            if isinstance(exc, RuntimeError) and "Gemini devolvió un error temporal" in str(exc):
                                st.error(
                                    "El Agente 2 no pudo completar la generación porque Gemini respondió temporalmente con saturación. "
                                    "Vuelve a intentarlo en unos minutos."
                                )
                            else:
                                st.error(f"Ocurri? un error en el procesamiento combinado: {exc}")
else:
    st.info("Ingresa la ruta de la carpeta del cliente para comenzar.")
